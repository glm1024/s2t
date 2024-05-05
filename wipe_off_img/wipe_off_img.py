import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
import threading
from tkinter import messagebox

# 定义要遍历的文件夹路径
root_folder = ''  # 初始化为空字符串

# 定义一个函数以打开文件对话框并获取文件夹路径
def get_folder():
    global root_folder
    folder_path = filedialog.askdirectory()
    if folder_path:
        root_folder = folder_path
        print("所选文件夹：", root_folder)
        if (root_folder):
            # 显示所选文件夹
            label_selected_folder.config(text=root_folder)
            label.config(text="")
            # 删除图片按钮状态设置为可用
            button_convert_wipe_off_img.config(state="normal")

# 定义一个函数以在新线程中删除Word文档中的所有图片
def remove_images_from_folder_thread():
    t = threading.Thread(target=remove_images_from_folder)
    t.start()

# 定义一个函数以遍历文件夹及子文件夹，并删除Word文档中的所有图片
def remove_images_from_folder():
    # 弹窗提醒和确认
    user_confirmation = messagebox.askyesno("确认操作", 
                                            "该操作会遍历所选文件夹及其子文件夹，并删除掉docx文件中的所有图片，处理后的文件会覆盖原文件，请确认是否继续？",
                                            parent=root)
    if not user_confirmation:
        return;

    # 按钮状态设置为不可用
    button_convert_wipe_off_img.config(state="disabled") 

    # 使用os.walk()遍历文件夹及子文件夹
    for folder_name, subfolders, filenames in os.walk(root_folder):
        for filename in filenames:
            # 检查文件名是否以.docx结尾
            if filename.endswith('.docx'):
                # 构建docx文件的完整路径
                docx_file_path = os.path.join(folder_name, filename)
                # 读取docx文件
                doc = Document(docx_file_path)
                
                # 遍历文档中的所有图片段落
                image_paragraphs = []
                for i in range(len(doc.paragraphs)):
                    images = doc.paragraphs[i]._element.xpath('.//pic:pic')
                    if len(images) > 0:
                        image_paragraphs.append(i)

                # 翻转段落顺序，从后往前删除图片，避免删除图片后导致段落索引变化
                image_paragraphs.reverse()

                # 删除图片段落
                for i in image_paragraphs:
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
                    p._p = p._element = None

                # 保存修改后的文档
                doc.save(docx_file_path)
                print("从文档中删除图片：", docx_file_path)

                # 更新显示
                sub_filename = filename[:7]
                label_selected_folder.config(text=sub_filename + "..." )

    # 处理完成
    label_selected_folder.config(text="处理完成！")

# 创建主窗口
root = tk.Tk()
root.title("docx文件图片删除工具")

# 创建一个容器
container = tk.Frame(root)
container.pack(pady=260)

# 创建一个标签
label = tk.Label(container, text="请选择一个文件夹：")
label.pack(side="left")  # 设置标签在左侧显示

# 创建一个标签以显示所选文件夹
label_selected_folder = tk.Label(container, text="")
label_selected_folder.pack(side="left")  # 设置标签在左侧显示

# 创建一个按钮以打开文件对话框
button = tk.Button(container, text="浏览", command=get_folder)
button.pack(side="left")  # 设置按钮在左侧显示，并添加水平间距

# 清空图片按钮
button_convert_wipe_off_img = tk.Button(container, text="删除图片", command=remove_images_from_folder_thread)
button_convert_wipe_off_img.pack(side="left", padx=10)  # 设置按钮在左侧显示，并添加水平间距
button_convert_wipe_off_img.config(state="disabled")  # 按钮默认状态为禁用

# 主窗口屏幕正中显示
# 计算屏幕尺寸
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 计算窗口尺寸
window_width = 800
window_height = 600

# 计算窗口位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口位置
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 禁止调整窗口大小
root.resizable(width=False, height=False) 

# 运行 GUI 事件循环
root.mainloop()